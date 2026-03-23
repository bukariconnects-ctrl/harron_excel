"""
Arabic Inventory Item Card Generator
Creates Excel (Item_Card.xlsx), PDF (Item_Card.pdf), and Word (Item_Card.docx) files with RTL layout for Arabic text
Requires: pip install openpyxl reportlab arabic-reshaper python-bidi python-docx
"""

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_item_card():
    # Create a new workbook and select the active sheet
    wb = Workbook()
    ws = wb.active
    ws.title = "كارت صنف"
    
    # Set sheet direction to RTL (Right-to-Left)
    ws.sheet_view.rightToLeft = True
    
    # Set default font
    default_font = Font(name='Arial', size=12)
    
    # Define border styles
    thick_border = Border(
        left=Side(style='medium', color='000000'),
        right=Side(style='medium', color='000000'),
        top=Side(style='medium', color='000000'),
        bottom=Side(style='medium', color='000000')
    )
    
    thin_border = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )
    
    dotted_horizontal = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='dotted', color='000000'),
        bottom=Side(style='dotted', color='000000')
    )
    
    # Set column widths
    column_widths = {
        'A': 12,  # التاريخ
        'B': 12,  # رقم الإذن
        'C': 30,  # البيان
        'D': 12,  # الوارد
        'E': 12,  # المنصرف
        'F': 12,  # الرصيد
        'G': 15   # ملاحظات
    }
    
    for col, width in column_widths.items():
        ws.column_dimensions[col].width = width
    
    # TOP SECTION
    # D2: Title "كارت صنف"
    ws['D2'] = 'كارت صنف'
    ws['D2'].font = Font(name='Arial', size=16, bold=True)
    ws['D2'].alignment = Alignment(horizontal='center', vertical='center')
    
    # A3: "مسلسل :"
    ws['A3'] = 'مسلسل :'
    ws['A3'].font = Font(name='Arial', size=12, bold=True)
    ws['A3'].alignment = Alignment(horizontal='right', vertical='center')
    
    # A4: "اسم الصنف :"
    ws['A4'] = 'اسم الصنف :'
    ws['A4'].font = Font(name='Arial', size=12, bold=True)
    ws['A4'].alignment = Alignment(horizontal='right', vertical='center')
    
    # D4: "رقم الصنف :"
    ws['D4'] = 'رقم الصنف :'
    ws['D4'].font = Font(name='Arial', size=12, bold=True)
    ws['D4'].alignment = Alignment(horizontal='center', vertical='center')
    
    # G4: "الوحدة :"
    ws['G4'] = 'الوحدة :'
    ws['G4'].font = Font(name='Arial', size=12, bold=True)
    ws['G4'].alignment = Alignment(horizontal='left', vertical='center')
    
    # TABLE HEADERS
    # Single row headers (Row 6)
    ws['A6'] = 'التاريخ'
    ws['A6'].font = Font(name='Arial', size=12, bold=True)
    ws['A6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A6'].border = thick_border
    
    ws['B6'] = 'رقم الإذن'
    ws['B6'].font = Font(name='Arial', size=12, bold=True)
    ws['B6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['B6'].border = thick_border
    
    ws['C6'] = 'البيان'
    ws['C6'].font = Font(name='Arial', size=12, bold=True)
    ws['C6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['C6'].border = thick_border
    
    ws['D6'] = 'الوارد'
    ws['D6'].font = Font(name='Arial', size=12, bold=True)
    ws['D6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['D6'].border = thick_border
    
    ws['E6'] = 'المنصرف'
    ws['E6'].font = Font(name='Arial', size=12, bold=True)
    ws['E6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['E6'].border = thick_border
    
    ws['F6'] = 'الرصيد'
    ws['F6'].font = Font(name='Arial', size=12, bold=True)
    ws['F6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['F6'].border = thick_border
    
    ws['G6'] = 'ملاحظات'
    ws['G6'].font = Font(name='Arial', size=12, bold=True)
    ws['G6'].alignment = Alignment(horizontal='center', vertical='center')
    ws['G6'].border = thick_border
    
    # DATA ROWS (Rows 7 to 31 - 25 empty rows)
    for row in range(7, 32):
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            cell = ws[f'{col}{row}']
            cell.font = default_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Apply borders with dotted horizontal lines
            cell.border = Border(
                left=Side(style='thin', color='000000'),
                right=Side(style='thin', color='000000'),
                top=Side(style='dotted', color='000000'),
                bottom=Side(style='dotted', color='000000')
            )
    
    # Apply solid bottom border to the last row
    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        cell = ws[f'{col}31']
        cell.border = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style='dotted', color='000000'),
            bottom=Side(style='medium', color='000000')
        )
    
    # Set row heights for better appearance
    ws.row_dimensions[2].height = 25
    ws.row_dimensions[6].height = 20
    
    # Save the workbook
    wb.save('Item_Card.xlsx')
    print("✓ Item_Card.xlsx has been created successfully!")
    print("✓ The file contains an Arabic inventory item card with RTL layout")

def reshape_arabic(text):
    """Helper function to reshape Arabic text for proper display"""
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text

def create_item_card_pdf():
    """Create PDF version of the item card with Arabic RTL support using Canvas"""
    
    pdf_filename = 'Item_Card.pdf'
    c = canvas.Canvas(pdf_filename, pagesize=A4)
    width, height = A4
    
    # Set font (using built-in Helvetica which supports basic Arabic when reshaped)
    c.setFont("Helvetica-Bold", 16)
    
    # Title
    title_text = reshape_arabic('كارت صنف')
    c.drawCentredString(width/2, height - 50, title_text)
    
    # Top section
    c.setFont("Helvetica-Bold", 10)
    y_pos = height - 90
    
    # Row 1: مسلسل, رقم الصنف, الوحدة
    c.drawRightString(width - 50, y_pos, reshape_arabic('مسلسل :'))
    c.drawCentredString(width/2, y_pos, reshape_arabic('رقم الصنف :'))
    c.drawString(50, y_pos, reshape_arabic('الوحدة :'))
    
    # Row 2: اسم الصنف
    y_pos -= 20
    c.drawRightString(width - 50, y_pos, reshape_arabic('اسم الصنف :'))
    
    # Table starting position
    y_pos -= 40
    table_start_y = y_pos
    
    # Table dimensions
    margin = 50
    table_width = width - 2 * margin
    col_widths = [table_width/7] * 7  # 7 equal columns
    row_height = 20
    
    # Draw table headers
    c.setFont("Helvetica-Bold", 10)
    headers = ['التاريخ', 'رقم الإذن', 'البيان', 'الوارد', 'المنصرف', 'الرصيد', 'ملاحظات']
    x_pos = width - margin
    
    for i, header in enumerate(headers):
        # Draw header cell border (thick)
        c.setLineWidth(1.5)
        c.rect(x_pos - col_widths[i], y_pos - row_height, col_widths[i], row_height)
        
        # Draw header text
        text = reshape_arabic(header)
        c.drawCentredString(x_pos - col_widths[i]/2, y_pos - row_height + 5, text)
        x_pos -= col_widths[i]
    
    # Draw data rows (25 rows)
    c.setFont("Helvetica", 9)
    c.setLineWidth(0.5)
    y_pos -= row_height
    
    for row in range(25):
        x_pos = width - margin
        for col in range(7):
            # Draw cell border
            c.rect(x_pos - col_widths[col], y_pos - row_height, col_widths[col], row_height)
            x_pos -= col_widths[col]
        y_pos -= row_height
        
        # Stop if we're running out of page space
        if y_pos < 100:
            break
    
    c.save()
    print("✓ Item_Card.pdf has been created successfully!")
    print("✓ The PDF contains an Arabic inventory item card with RTL layout")

def create_item_card_word():
    """Create Word document version of the item card with Arabic RTL support"""
    
    doc = Document()
    
    # Set document to RTL
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)
    
    # Title
    title = doc.add_paragraph('كارت صنف')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    
    # Top section
    doc.add_paragraph()
    top_para1 = doc.add_paragraph()
    top_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = top_para1.add_run('مسلسل :                    رقم الصنف :                    الوحدة :')
    run1.font.bold = True
    run1.font.size = Pt(12)
    run1.font.name = 'Arial'
    
    top_para2 = doc.add_paragraph()
    top_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = top_para2.add_run('اسم الصنف :')
    run2.font.bold = True
    run2.font.size = Pt(12)
    run2.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # Create table (1 header row + 25 data rows, 7 columns)
    table = doc.add_table(rows=26, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set RTL for table
    tbl = table._element
    tblPr = tbl.tblPr
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)
    
    # Headers (RTL order)
    headers = ['التاريخ', 'رقم الإذن', 'البيان', 'الوارد', 'المنصرف', 'الرصيد', 'ملاحظات']
    header_cells = table.rows[0].cells
    
    for i, header_text in enumerate(headers):
        cell = header_cells[6-i]  # Reverse order for RTL
        cell.text = header_text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
    
    # Set column widths
    col_widths_inches = [0.9, 0.9, 2.5, 0.9, 0.9, 0.9, 1.2]  # Adjusted for 7 columns
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths_inches[i])
            # Center align all cells
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    # Save document
    doc.save('Item_Card.docx')
    print("✓ Item_Card.docx has been created successfully!")
    print("✓ The Word document contains an Arabic inventory item card with RTL layout")

if __name__ == "__main__":
    create_item_card()
    create_item_card_pdf()
    create_item_card_word()
